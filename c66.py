import streamlit as st
import chromadb
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer
from google.generativeai import GenerativeModel, configure
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import io
import os
import csv
import pandas as pd
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
import time
import base64
import json
import plotly.express as px
import plotly.graph_objects as go
from docx import Document

# --------------------
# Page Configuration
# --------------------
st.set_page_config(
    page_title="🤖 AI Job Matcher & Scheduler",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --------------------
# Session State Reset on Refresh
# --------------------
def initialize_session_state():
    """Initialize or reset session state variables"""
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        st.session_state.step1_done = False
        st.session_state.step2_done = False
        st.session_state.jd_text = ""
        st.session_state.selected_candidate = None
        st.session_state.candidate_email = None
        st.session_state.last_rating_result = None
        st.session_state.questions_generated = {}
        st.session_state.scheduled_interviews = {}  # ✅ Track scheduled interviews

initialize_session_state()

# --------------------
# Secure Gemini / LLM setup
# --------------------
configure(api_key="AIzaSyAsoHAC3qKwCbcpMauiUI8kzSLfcZPZAyQ")
llm_model = GenerativeModel("gemini-2.5-flash-lite")

# --------------------
# Chroma + Embedding Model
# --------------------
@st.cache_resource
def get_chroma_client():
    client = chromadb.Client()
    collection = client.get_or_create_collection("resumes")
    return client, collection

@st.cache_resource
def get_embedder():
    return SentenceTransformer("all-MiniLM-L6-v2")

chroma_client, collection = get_chroma_client()
embedder = get_embedder()

# --------------------
# Constants for email
# --------------------
SENDER_EMAIL = "gdhivya2302@gmail.com"
SENDER_PASSWORD = "jaho mjob oepq rtnu"
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587

# --------------------
# ✅ Helper function to clean and normalize recommendation values
# --------------------
def clean_recommendation(value):
    """Clean recommendation value to ensure exact match"""
    if pd.isna(value):
        return "Review"
    
    import re
    cleaned = str(value).strip()
    cleaned = re.sub(r'[*_`\[\]"\']', '', cleaned)
    cleaned = cleaned.strip()
    
    if "hire" in cleaned.lower():
        return "Hire"
    elif "reject" in cleaned.lower():
        return "Reject"
    elif "review" in cleaned.lower():
        return "Review"
    else:
        return "Review"

def clean_sentiment(value):
    """Clean sentiment value to ensure exact match"""
    if pd.isna(value):
        return "Neutral"
    
    import re
    cleaned = str(value).strip()
    cleaned = re.sub(r'[*_`\[\]"\']', '', cleaned)
    cleaned = cleaned.strip()
    
    if "positive" in cleaned.lower():
        return "Positive"
    elif "negative" in cleaned.lower():
        return "Negative"
    elif "neutral" in cleaned.lower():
        return "Neutral"
    else:
        return "Neutral"

# --------------------
# Helper: Extract Text from PDF
# --------------------
def extract_text_from_pdf(file_obj):
    reader = PdfReader(file_obj)
    text = ""
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text += f"\n[Page {i+1}]\n{page_text.strip()}\n"
    return text.strip()

# --------------------
# Store Resume in ChromaDB
# --------------------
def store_resume_text(text, filename):
    if not text:
        return "❌ No text extracted from PDF."
    embedding = embedder.encode(text).tolist()
    try:
        collection.add(
            documents=[text],
            embeddings=[embedding],
            ids=[filename],
            metadatas=[{"filename": filename}]
        )
        return f"✅ Stored resume: {filename}"
    except Exception as e:
        return f"⚠️ Could not store {filename}: {e}"

# --------------------
# Google Drive Functions
# --------------------
DRIVE_SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
CALENDAR_SCOPES = ['https://www.googleapis.com/auth/calendar']

def connect_drive():
    flow = InstalledAppFlow.from_client_secrets_file("credentials.json", DRIVE_SCOPES)
    creds = flow.run_local_server(port=0)
    service = build('drive', 'v3', credentials=creds)
    return service

def connect_calendar():
    """Connect to Google Calendar API"""
    flow = InstalledAppFlow.from_client_secrets_file("credentials.json", CALENDAR_SCOPES)
    creds = flow.run_local_server(port=0)
    service = build('calendar', 'v3', credentials=creds)
    return service

def fetch_resumes_from_drive(service, folder_id):
    results = service.files().list(
        q=f"'{folder_id}' in parents and mimeType='application/pdf'",
        fields="files(id, name)"
    ).execute()
    files = results.get("files", [])
    downloaded = []
    for file in files:
        file_id = file["id"]
        file_name = file["name"]
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        text = extract_text_from_pdf(fh)
        if text:
            store_msg = store_resume_text(text, file_name)
            downloaded.append({"filename": file_name, "text": text, "store_msg": store_msg})
        else:
            downloaded.append({"filename": file_name, "text": "", "store_msg": "No text extracted."})
    return downloaded

# --------------------
# Candidate Details Extraction
# --------------------
def extract_candidate_details(resume_text):
    prompt = f"""
Extract the following details from this resume text:
1. Full Name
2. Email Address (if present)

Resume Text:
{resume_text}

Output format exactly:
Name: [Full Name]
Email: [Email Address or Not provided]
"""
    try:
        response = llm_model.generate_content(prompt)
        output = response.text.strip()
        name, email = "Unknown", "Not provided"
        if "Name:" in output:
            name = output.split("Name:")[1].split("\n")[0].strip()
        if "Email:" in output:
            email = output.split("Email:")[1].split("\n")[0].strip()
        return {"name": name, "email": email}
    except Exception as e:
        st.error(f"Error extracting details: {e}")
        return {"name": "Unknown", "email": "Not provided"}

# --------------------
# ✅ Interview Questions Generator
# --------------------
def generate_interview_questions(job_description, resume_text, candidate_name):
    """Generate personalized interview questions and save as Word document"""
    prompt = f"""
You are an expert HR interviewer with 15+ years of experience. Based on the job description 
and candidate's resume, generate 10-15 highly relevant, insightful interview questions.

Your questions should assess:
1. **Technical Skills** (3-4 questions) - Role-specific technical competencies
2. **Experience Relevance** (2-3 questions) - Past work experience alignment
3. **Problem-Solving Ability** (2-3 questions) - Analytical and critical thinking
4. **Cultural Fit** (2-3 questions) - Values, work style, team dynamics
5. **Motivation & Goals** (2-3 questions) - Career aspirations, interest in role

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 JOB DESCRIPTION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{job_description}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📄 CANDIDATE RESUME
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{resume_text}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ INSTRUCTIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Generate questions that are:
✓ Specific to this candidate's background (reference their actual experience)
✓ Probing and thought-provoking (avoid yes/no questions)
✓ Role-appropriate (match job requirements)
✓ Behavioral when relevant (use STAR format prompts)
✓ Open-ended to encourage detailed responses

Output format (strictly follow):

TECHNICAL SKILLS
1. [Question about specific technical skill from JD]
2. [Question about technical problem-solving]
3. [Question about relevant tools/technologies]

EXPERIENCE RELEVANCE
4. [Question about specific past project mentioned in resume]
5. [Question about handling similar responsibilities]

PROBLEM-SOLVING ABILITY
6. [Situational question related to role challenges]
7. [Question about overcoming technical obstacles]

CULTURAL FIT
8. [Question about work style and collaboration]
9. [Question about handling feedback/conflict]

MOTIVATION & GOALS
10. [Question about interest in this specific role]
11. [Question about career aspirations]
12. [Question about what excites them about the company]

[Continue with 3-4 more questions as needed, ensuring diversity across categories]

Remember: Personalize questions using actual details from the candidate's resume!
"""
    try:
        response = llm_model.generate_content(prompt)
        questions_text = response.text.strip()
        
        # Create Word document
        doc = Document()
        
        # Title
        doc.add_heading(f'Interview Questions for {candidate_name}', 0)
        
        # Metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run('Generated on: ').bold = True
        metadata_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        metadata_para.add_run('\nCandidate: ').bold = True
        metadata_para.add_run(candidate_name)
        
        doc.add_paragraph('')
        
        # Job Description Section
        doc.add_heading('📋 Job Description', level=1)
        doc.add_paragraph(job_description)
        
        doc.add_paragraph('')
        
        # Interview Questions Section
        doc.add_heading('❓ Interview Questions', level=1)
        
        questions_paragraphs = questions_text.split('\n')
        for line in questions_paragraphs:
            line = line.strip()
            if not line:
                continue
            
            if line.isupper() or any(keyword in line.upper() for keyword in 
                                    ['TECHNICAL', 'EXPERIENCE', 'PROBLEM', 'CULTURAL', 'MOTIVATION']):
                doc.add_heading(line, level=2)
            else:
                doc.add_paragraph(line, style='List Number' if line[0].isdigit() else 'Normal')
        
        # Additional Notes Section
        doc.add_paragraph('')
        doc.add_heading('📝 Interviewer Notes', level=1)
        doc.add_paragraph('Use this space to record candidate responses and observations during the interview:')
        doc.add_paragraph('\n' * 15)
        
        # Save document
        os.makedirs("interview_questions", exist_ok=True)
        filename = f"{candidate_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_Interview_Questions.docx"
        filepath = os.path.join("interview_questions", filename)
        doc.save(filepath)
        
        return filename, filepath, questions_text
        
    except Exception as e:
        st.error(f"Error generating questions: {e}")
        return None, None, f"Error: {e}"

# --------------------
# CSV Handling
# --------------------
def clear_matched_candidates(file_name="matched_candidates.csv"):
    if os.path.exists(file_name):
        os.remove(file_name)

def store_candidate_data(candidate_details, file_name="matched_candidates.csv"):
    file_exists = os.path.exists(file_name)
    with open(file_name, mode='a', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=[
            "name", "email", "cosine_score", "llm_confidence", 
            "resume_text", "filename", "disclaimer", "interview_questions_doc"
        ])

        if not file_exists:
            writer.writeheader()
        writer.writerow(candidate_details)

# --------------------
# Interview Ratings CSV
# --------------------
def store_interview_rating(rating_data, file_name="interview_ratings.csv"):
    """Store interview rating with cleaned recommendation and sentiment values"""
    rating_data["recommendation"] = clean_recommendation(rating_data.get("recommendation", "Review"))
    rating_data["sentiment"] = clean_sentiment(rating_data.get("sentiment", "Neutral"))
    
    file_exists = os.path.exists(file_name)
    with open(file_name, mode='a', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=[
            "name", "email", "technical_rating", "communication_rating", 
            "coding_rating", "problem_solving_rating",
            "interviewer_notes", "sentiment", "recommendation", 
            "llm_analysis", "timestamp"
        ])
        if not file_exists:
            writer.writeheader()
        writer.writerow(rating_data)

# --------------------
# IMPROVED LLM Feedback Analysis
# --------------------
def analyze_interview_feedback(technical, communication, coding, problem_solving, notes):
    """Enhanced LLM-based interview feedback analysis with comprehensive evaluation criteria"""
    import re
    
    avg_score = (technical + communication + coding + problem_solving) / 4
    
    prompt = f"""
You are a senior HR analyst and talent acquisition expert with 15+ years of experience in candidate evaluation. 
Conduct a comprehensive, professional analysis of this interview feedback.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 INTERVIEW PERFORMANCE METRICS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Quantitative Ratings (Scale: 1-10):
- Technical Skills:     {technical}/10
- Communication:        {communication}/10  
- Coding Ability:       {coding}/10
- Problem Solving:      {problem_solving}/10
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- Overall Average:      {avg_score:.2f}/10
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 Qualitative Assessment:
{notes if notes.strip() else "No additional notes provided by interviewer."}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 REQUIRED ANALYSIS COMPONENTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Provide a thorough, data-driven evaluation with the following structure:

1. **SENTIMENT ANALYSIS**
   Classify the overall interview sentiment based on both scores and notes:
   
   • Positive: Strong performance (avg ≥7.0) with encouraging feedback and demonstrated competencies
   • Neutral: Average performance (avg 5.0-6.9) with mixed indicators or areas needing clarification  
   • Negative: Below expectations (avg <5.0) with significant gaps or concerning patterns
   
   Consider: Score patterns, consistency across dimensions, tone of interviewer notes

2. **HIRING RECOMMENDATION**
   Provide ONE clear recommendation with solid justification:
   
   • Hire: Candidate meets or exceeds requirements
     - Criteria: avg ≥7.0, strong performance in critical areas, positive interviewer sentiment
     - Confidence level: High
   
   • Review: Candidate shows potential but requires further evaluation
     - Criteria: avg 5.5-6.9, uneven performance, or insufficient data points
     - Suggested next steps: Additional technical round, skills assessment, or reference checks
   
   • Reject: Candidate does not meet minimum standards
     - Criteria: avg <5.5, critical skill gaps, or significant concerns raised
     - Be constructive but clear about decision rationale

3. **COMPREHENSIVE PROFESSIONAL ANALYSIS**
   Write a detailed 5-7 sentence evaluation that covers:

   ✓ **Performance Overview**: 
     - Synthesize quantitative scores into narrative context
     - Highlight overall impression and readiness for the role
   
   ✓ **Key Strengths** (be specific):
     - Identify 2-3 standout competencies with evidence
     - Reference specific high ratings (8+) and positive notes
     - Connect strengths to role requirements
   
   ✓ **Areas of Concern** (if any):
     - Note any ratings below 6 or red flags from notes
     - Be constructive: frame as "development areas" not just weaknesses
     - Assess whether concerns are deal-breakers or coachable
   
   ✓ **Soft Skills & Cultural Fit**:
     - Evaluate communication quality, confidence, problem-solving approach
     - Consider collaboration signals, adaptability, learning mindset
   
   ✓ **Decision Justification**:
     - Explicitly connect your recommendation to the evidence
     - Address any conflicting signals (e.g., high tech score but low communication)
     - Explain how the candidate compares to typical role expectations
   
   ✓ **Actionable Next Steps**:
     - If Hire: Suggest offer timeline, onboarding focus areas
     - If Review: Specify what additional data is needed (e.g., "schedule pair programming session")
     - If Reject: Note if candidate might be suitable for different role

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ CRITICAL OUTPUT FORMAT REQUIREMENTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

IMPORTANT: You MUST format your response EXACTLY as shown below. Do NOT use markdown formatting, asterisks, or any special characters in the labels.

SENTIMENT: [ONLY write one word: Positive OR Neutral OR Negative]

RECOMMENDATION: [ONLY write one word: Hire OR Review OR Reject]

ANALYSIS:
[Write your comprehensive 5-7 sentence professional evaluation here. This should read like a formal HR assessment that could be presented to hiring managers. Be thorough, specific, and reference the actual data points. Structure your analysis to flow logically through the components outlined above: performance overview → strengths → concerns → soft skills → decision justification → next steps.]

EXAMPLE OUTPUT:
SENTIMENT: Positive

RECOMMENDATION: Hire

ANALYSIS:
The candidate demonstrated exceptional technical proficiency with an average score of 8.5/10, indicating strong readiness for the role. Their technical skills rating of 9/10 and coding ability of 8/10 showcase solid foundational competencies in required technologies. Communication was rated at 8/10, suggesting they can effectively collaborate with cross-functional teams and articulate complex technical concepts. Problem-solving abilities scored 9/10, reflecting critical thinking and analytical strength essential for this position. The interviewer's notes highlighted their enthusiasm and cultural alignment with company values. Based on the consistently high performance across all dimensions and positive sentiment, this candidate is recommended for hire with high confidence. Next steps should include extending an offer within 48 hours to secure this strong candidate before competing offers emerge.

Remember: Your analysis directly influences hiring decisions. Be thorough, fair, and professional.
"""
    
    try:
        response = llm_model.generate_content(prompt)
        output = response.text.strip()
        
        sentiment = "Neutral"
        recommendation = "Review"
        analysis = "Analysis not available due to parsing error."
        
        # Robust parsing with regex
        sentiment_match = re.search(r'SENTIMENT:\s*\*?\*?(\w+)\*?\*?', output, re.IGNORECASE)
        recommendation_match = re.search(r'RECOMMENDATION:\s*\*?\*?(\w+)\*?\*?', output, re.IGNORECASE)
        analysis_match = re.search(r'ANALYSIS:\s*\n(.+)', output, re.DOTALL | re.IGNORECASE)
        
        if sentiment_match:
            extracted_sentiment = sentiment_match.group(1).strip()
            extracted_sentiment = re.sub(r'[*_`\[\]"\']', '', extracted_sentiment).strip()
            for s in ["Positive", "Neutral", "Negative"]:
                if s.lower() in extracted_sentiment.lower():
                    sentiment = s
                    break
        
        if recommendation_match:
            extracted_recommendation = recommendation_match.group(1).strip()
            extracted_recommendation = re.sub(r'[*_`\[\]"\']', '', extracted_recommendation).strip()
            for r in ["Hire", "Review", "Reject"]:
                if r.lower() in extracted_recommendation.lower():
                    recommendation = r
                    break
        
        if analysis_match:
            analysis = analysis_match.group(1).strip()
            analysis = re.sub(r'```[\w]*\n?', '', analysis)
            analysis = re.sub(r'\*\*([^*]+)\*\*', r'\1', analysis)
            analysis = re.sub(r'\*([^*]+)\*', r'\1', analysis)
            analysis = analysis.strip()
        
        # Fallback parsing
        if sentiment == "Neutral" and "SENTIMENT:" in output.upper():
            try:
                sentiment_line = output.split("SENTIMENT:")[1].split("\n")[0].strip()
                sentiment_line = sentiment_line.replace("*", "").replace('"', '').replace("'", "").strip()
                for s in ["Positive", "Neutral", "Negative"]:
                    if s.lower() in sentiment_line.lower():
                        sentiment = s
                        break
            except:
                pass
        
        if recommendation == "Review" and "RECOMMENDATION:" in output.upper():
            try:
                rec_line = output.split("RECOMMENDATION:")[1].split("\n")[0].strip()
                rec_line = rec_line.replace("*", "").replace('"', '').replace("'", "").strip()
                for r in ["Hire", "Review", "Reject"]:
                    if r.lower() in rec_line.lower():
                        recommendation = r
                        break
            except:
                pass
        
        # Score-based validation
        if recommendation == "Review":
            if avg_score >= 7.5:
                recommendation = "Hire"
                st.warning("⚠️ Auto-corrected recommendation to 'Hire' based on high average score.")
            elif avg_score < 5.0:
                recommendation = "Reject"
                st.warning("⚠️ Auto-corrected recommendation to 'Reject' based on low average score.")
        
        sentiment = clean_sentiment(sentiment)
        recommendation = clean_recommendation(recommendation)
        
        return {
            "sentiment": sentiment,
            "recommendation": recommendation,
            "analysis": analysis,
            "avg_score": avg_score
        }
        
    except Exception as e:
        st.error(f"⚠️ LLM Analysis Error: {str(e)}")
        
        if avg_score >= 7.5:
            sentiment = "Positive"
            recommendation = "Hire"
        elif avg_score >= 5.5:
            sentiment = "Neutral"
            recommendation = "Review"
        else:
            sentiment = "Negative"
            recommendation = "Reject"
        
        return {
            "sentiment": sentiment,
            "recommendation": recommendation,
            "analysis": f"Automated analysis based on scores: The candidate achieved an average rating of {avg_score:.2f}/10 (Technical: {technical}/10, Communication: {communication}/10, Coding: {coding}/10, Problem Solving: {problem_solving}/10). {notes[:200] if notes else 'No additional notes provided.'}",
            "avg_score": avg_score
        }

# --------------------
# Utility: Clear Resume DB
# --------------------
def clear_resume_database():
    global collection, chroma_client
    try:
        chroma_client.delete_collection("resumes")
        collection = chroma_client.get_or_create_collection("resumes")
        clear_matched_candidates()
        return "✅ Database and CSV cleared successfully!"
    except Exception as e:
        return f"⚠️ Error clearing database: {e}"

# --------------------
# Matching Logic
# --------------------
def find_best_match_with_llm(job_description, percentage=50):
    if collection.count() == 0:
        return "⚠️ No resumes uploaded yet.", []

    clear_matched_candidates()
    all_data = collection.get(include=["metadatas", "documents", "embeddings"])
    similarities = []

    jd_embedding = embedder.encode(job_description).reshape(1, -1)

    for i, emb in enumerate(all_data["embeddings"]):
        cos_score = cosine_similarity(jd_embedding, np.array(emb).reshape(1, -1))[0][0]

        prompt = f"""
You are a recruitment AI. Analyze this candidate for the following job description
and give a confidence score between 0 and 1.

Job Description:
{job_description}

Resume:
{all_data['documents'][i]}

Output format: Confidence: 0.85
"""
        try:
            response = llm_model.generate_content(prompt)
            llm_conf = 0.0
            if "Confidence:" in response.text:
                try:
                    llm_conf = float(response.text.split("Confidence:")[1].split()[0])
                except:
                    llm_conf = 0.0
        except Exception:
            llm_conf = 0.0

        similarities.append({
            "index": i,
            "cosine_score": cos_score,
            "llm_confidence": llm_conf,
            "filename": all_data["metadatas"][i]["filename"],
            "content": all_data["documents"][i]
        })

    for s in similarities:
        s["combined_score"] = (s["cosine_score"] + s["llm_confidence"]) / 2

    similarities.sort(key=lambda x: x["combined_score"], reverse=True)

    top_n = max(1, int(len(similarities) * (percentage / 100)))
    top_candidates = similarities[:top_n]

    candidate_results = []
    for c in top_candidates:
        details = extract_candidate_details(c["content"])

        disclaimer_prompt = f"""
Analyze this candidate's resume and the given job description.
Provide a short disclaimer summarizing their key strengths (pros) and weaknesses or gaps (cons)
relevant to the role. Be concise — use bullet points under each section.

Job Description:
{job_description}

Candidate Resume:
{c["content"]}

Output format:
Pros:
- ...
- ...

Cons:
- ...
- ...
"""
        try:
            disclaimer_resp = llm_model.generate_content(disclaimer_prompt)
            disclaimer_text = disclaimer_resp.text.strip()
        except Exception:
            disclaimer_text = "Disclaimer not available (LLM error)."

        details.update({
            "cosine_score": round(c["cosine_score"], 3),
            "llm_confidence": round(c["llm_confidence"], 3),
            "resume_text": c["content"],
            "filename": c["filename"],
            "disclaimer": disclaimer_text,
            "interview_questions_doc": "Not generated"
        })

        store_candidate_data(details)
        candidate_results.append(details)

    if candidate_results:
        best = candidate_results[0]
        prompt = f"""
You are an expert recruiter AI. Analyze the following job description and candidate profile carefully.

Step 1️⃣: Check whether the job description requires experience.
Step 2️⃣: Check if the candidate has mentioned experience.  
If the job requires experience and the candidate does NOT have it, clearly mention this as a mismatch reason first.  
Then check for other skills that may compensate.

Finally, summarize:
- Experience Match: Yes/No
- Skill Match: High/Medium/Low
- Reasoning: Explain briefly why or why not the candidate fits.
- Final Verdict: Is this candidate a good fit overall?

Job Description:
{job_description}

Candidate Resume:
{best['resume_text']}

Candidate Name: {best['name']}
"""
        try:
            response = llm_model.generate_content(prompt)
            explanation = response.text.strip()
            return explanation, candidate_results
        except Exception:
            return f"Top score: {best['cosine_score']}", candidate_results

    return "No suitable candidates found.", []

# --------------------
# Email Logic
# --------------------
def send_email(to_email, subject, body):
    try:
        msg = MIMEMultipart()
        msg["From"] = SENDER_EMAIL
        msg["To"] = to_email
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain"))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SENDER_EMAIL, SENDER_PASSWORD)
            server.send_message(msg)
        return f"✅ Email sent to {to_email}"
    except Exception as e:
        return f"⚠️ Failed: {e}"

def send_shortlist_email(name, email, job_description=None):
    jd = job_description if job_description else st.session_state.get("jd_text", None)

    # Generate next 3 days
    today = datetime.now()
    day1 = (today + timedelta(days=1)).strftime('%Y-%m-%d')
    day2 = (today + timedelta(days=2)).strftime('%Y-%m-%d')
    day3 = (today + timedelta(days=3)).strftime('%Y-%m-%d')
    
    subject = "You are Shortlisted for Interview"
    body = f"""
Dear {name},

🎉 Congratulations! You have been shortlisted for the interview process with our team.

Below is the job description for your reference:

📄 Job Description:
{jd if jd else "Job description not available."}

Kindly reply to this email with your available date and time for the interview. Please respond within the next 12 hours so we can schedule your slot accordingly.

Available Dates:
- {day1}
- {day2}
- {day3}

Example Reply:
I am available on {day1} at 5:30 PM.

Interview Details:
- Duration: 45–60 minutes
- Mode: Video Call (Google Meet link will be shared upon confirmation)

We look forward to connecting with you soon.

Best regards,
HR Team
"""
    return send_email(email, subject, body)
def create_google_calendar_link(name, interview_datetime, duration_minutes=60):
    """Generate a Google Calendar event link with pre-filled details"""
    from urllib.parse import quote
    from datetime import timedelta  # ensure timedelta is imported
    
    end_datetime = interview_datetime + timedelta(minutes=duration_minutes)
    
    start_str = interview_datetime.strftime('%Y%m%dT%H%M%S')
    end_str = end_datetime.strftime('%Y%m%dT%H%M%S')
    
    event_title = quote(f"Interview with HR Team - {name}")
    event_description = quote(
        f"Interview Session\n\n"
        f"Candidate: {name}\n"
        f"Duration: {duration_minutes} minutes\n"
        f"Mode: Video Call (Google Meet)\n\n"
        f"Meeting Link: https://meet.google.com/zag-iocp-kuz\n\n"
        f"Please join on time and keep your resume handy.\n"
        f"Ensure you have a stable internet connection."
    )
    event_location = quote("Google Meet: https://meet.google.com/zag-iocp-kuz")
    
    # ✅ Added Google Calendar notifications (reminders)
    calendar_url = (
        f"https://calendar.google.com/calendar/render?action=TEMPLATE"
        f"&text={event_title}"
        f"&dates={start_str}/{end_str}"
        f"&details={event_description}"
        f"&location={event_location}"
        f"&ctz=Asia/Kolkata"
        f"&add=notification"  # <-- this enables reminders/notifications
    )
    
    return calendar_url


def send_interview_email(name, email, interview_datetime, time_range_msg=None):
    subject = "Interview Scheduled - Invitation"
    
    time_display = interview_datetime.strftime('%d %B %Y, %I:%M %p IST')
    
    availability_note = ""
    if time_range_msg and "time range" in time_range_msg.lower():
        availability_note = "\n(Time selected based on your availability window)"
    
    calendar_link = create_google_calendar_link(name, interview_datetime)
    
    body = f"""
Dear {name},

Your interview has been scheduled as per your availability:

📅 Date & Time: {time_display}{availability_note}

Please join using the Google Meet link below at the scheduled time:
🔗 https://meet.google.com/zag-iocp-kuz 

Interview Details:
- Duration: 45–60 minutes
- Mode: Video Call (Google Meet)
- Please ensure you have a stable internet connection
- Keep your resume and relevant documents handy

We look forward to speaking with you!

Best regards,
HR Team
"""
    return send_email(email, subject, body)


# --------------------
# Gmail Reply Reader
# --------------------
GMAIL_SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']

def read_gmail_replies_and_extract_date(candidate_email):
    try:
        flow = InstalledAppFlow.from_client_secrets_file("credentials.json", GMAIL_SCOPES)
        creds = flow.run_local_server(port=0)
        service = build('gmail', 'v1', credentials=creds)

        query = f"from:{candidate_email}"
        res = service.users().messages().list(userId='me', q=query, maxResults=10).execute()
        msgs = res.get("messages", [])
        if not msgs:
            return None, "No messages found from candidate."

        for m in msgs:
            msg = service.users().messages().get(userId='me', id=m['id'], format='full').execute()
            snippet = msg.get('snippet', '')
            body_text = snippet
            payload = msg.get('payload', {})
            parts = payload.get('parts', []) or []
            for p in parts:
                if p.get('mimeType') == 'text/plain' and 'body' in p and 'data' in p['body']:
                    try:
                        data = p['body']['data']
                        text = base64.urlsafe_b64decode(data.encode('UTF-8')).decode('utf-8')
                        body_text = text
                    except Exception:
                        pass

            date_prompt = f"""
You are extracting interview availability from a candidate's email reply.

Task: Extract the date and time information from the candidate's response.

Instructions:
1. If they mention a specific date and time (e.g., "3rd November at 3:30 OR 3.00 PM"), output: DATE|TIME
2. If they mention a date with a time range (e.g., "3rd November between 3 PM to 5 PM"), output: DATE|START_TIME|END_TIME
3. If they mention only a date (no time), output: DATE|NONE
4. If no date is found, output: NONE

Output Format Examples:
- "2025-11-03|15:30" (specific date and time)
- "2025-11-03|15:00|17:00" (date with time range)
- "2025-11-03|NONE" (date only, no time)
- "NONE" (no date found)

Important:
- Use 24-hour format for times (HH:MM)
- Use YYYY-MM-DD format for dates
- Do not include any explanations, just the formatted output

Candidate reply:
\"\"\"{body_text}\"\"\""""
            
            resp = llm_model.generate_content(date_prompt)
            extracted = resp.text.strip().splitlines()[0].strip()
            
            if extracted.upper() == "NONE":
                continue
            
            try:
                parts = extracted.split("|")
                
                if len(parts) == 2:
                    date_str = parts[0]
                    time_str = parts[1]
                    
                    if time_str.upper() == "NONE":
                        dt = datetime.strptime(date_str, "%Y-%m-%d")
                        dt = datetime.combine(dt.date(), datetime.strptime("11:00", "%H:%M").time())
                        return dt, f"Parsed date-only: {date_str}, defaulted to 11:00 AM"
                    else:
                        dt = datetime.strptime(f"{date_str} {time_str}", "%Y-%m-%d %H:%M")
                        return dt, f"Parsed specific date+time: {date_str} {time_str}"
                
                elif len(parts) == 3:
                    date_str = parts[0]
                    start_time_str = parts[1]
                    end_time_str = parts[2]
                    
                    start_time = datetime.strptime(start_time_str, "%H:%M").time()
                    end_time = datetime.strptime(end_time_str, "%H:%M").time()
                    
                    start_minutes = start_time.hour * 60 + start_time.minute
                    end_minutes = end_time.hour * 60 + end_time.minute
                    mid_minutes = (start_minutes + end_minutes) // 2
                    
                    mid_hour = mid_minutes // 60
                    mid_minute = mid_minutes % 60
                    
                    dt = datetime.strptime(date_str, "%Y-%m-%d")
                    dt = datetime.combine(dt.date(), datetime.strptime(f"{mid_hour:02d}:{mid_minute:02d}", "%H:%M").time())
                    
                    return dt, f"Parsed date with time range: {date_str} from {start_time_str} to {end_time_str}, scheduled at {mid_hour:02d}:{mid_minute:02d}"
                
            except ValueError as ve:
                continue
                
        return None, "No valid date/time found."
    except Exception as e:
        return None, f"Error: {e}"

# --------------------
# Analytics Functions
# --------------------
def create_analytics_dashboard():
    """Create simplified analytics dashboard with 3 charts and 2-candidate comparison"""
    if not os.path.exists("interview_ratings.csv"):
        st.info("📊 No interview data available yet. Complete some interviews to see analytics.")
        return
    
    df = pd.read_csv("interview_ratings.csv")
    
    if len(df) == 0:
        st.info("📊 No interview data available yet.")
        return
    
    df['recommendation'] = df['recommendation'].apply(clean_recommendation)
    df['sentiment'] = df['sentiment'].apply(clean_sentiment)
    
    if 'timestamp' not in df.columns:
        df['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    df['avg_rating'] = df[["technical_rating", "communication_rating", "coding_rating", 
                           "problem_solving_rating"]].mean(axis=1)
    
    st.title("📊 Analytics Dashboard")
    
    # Filters
    st.markdown("### 🔍 Filters")
    col_filter1, col_filter2 = st.columns(2)
    
    with col_filter1:
        all_names = sorted(df["name"].unique().tolist())
        selected_names = st.multiselect("👤 Select Candidate(s):", all_names, default=[], key="name_filter")
    
    with col_filter2:
        all_sentiments = ["All Sentiments"] + sorted(df["sentiment"].unique().tolist())
        selected_sentiment = st.selectbox("😊 Select Sentiment:", all_sentiments, key="sentiment_filter")
    
    # Apply Filters
    filtered_df = df.copy()
    
    if selected_names:
        filtered_df = filtered_df[filtered_df["name"].isin(selected_names)]
    
    if selected_sentiment != "All Sentiments":
        filtered_df = filtered_df[filtered_df["sentiment"] == selected_sentiment]
    
    if len(filtered_df) == 0:
        st.warning("⚠️ No data matches the selected filters. Please adjust your selection.")
        return
    
    st.markdown("---")
    
    # Key Metrics
    st.markdown("### 📊 Key Metrics")
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        st.metric("Total Interviews", len(filtered_df))
    
    with col2:
        hire_count = len(filtered_df[filtered_df["recommendation"] == "Hire"])
        hire_pct = (hire_count/len(filtered_df)*100) if len(filtered_df) > 0 else 0
        st.metric("✅ Hire", hire_count, delta=f"{hire_pct:.1f}%")
    
    with col3:
        review_count = len(filtered_df[filtered_df["recommendation"] == "Review"])
        review_pct = (review_count/len(filtered_df)*100) if len(filtered_df) > 0 else 0
        st.metric("⚠️ Review", review_count, delta=f"{review_pct:.1f}%")
    
    with col4:
        reject_count = len(filtered_df[filtered_df["recommendation"] == "Reject"])
        reject_pct = (reject_count/len(filtered_df)*100) if len(filtered_df) > 0 else 0
        st.metric("❌ Reject", reject_count, delta=f"{reject_pct:.1f}%")
    
    with col5:
        avg_overall = filtered_df[["technical_rating", "communication_rating", "coding_rating", 
                                   "problem_solving_rating"]].mean().mean()
        st.metric("Avg Rating", f"{avg_overall:.1f}/10")
    
    st.markdown("---")
    
    # 2-Candidate Comparison
    if len(selected_names) == 2:
        st.markdown(f"## 🆚 Candidate Comparison: {selected_names[0]} vs {selected_names[1]}")
        
        candidate1_data = filtered_df[filtered_df["name"] == selected_names[0]].iloc[-1]
        candidate2_data = filtered_df[filtered_df["name"] == selected_names[1]].iloc[-1]
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown(f"#### 👤 {selected_names[0]}")
            st.write(f"**Email:** {candidate1_data['email']}")
            st.write(f"**Avg Rating:** {candidate1_data['avg_rating']:.2f}/10")
            st.write(f"**Sentiment:** {candidate1_data['sentiment']}")
            st.write(f"**Recommendation:** {candidate1_data['recommendation']}")
        
        with col2:
            st.markdown("#### ⚖️ Comparison")
            diff = candidate1_data['avg_rating'] - candidate2_data['avg_rating']
            if abs(diff) < 0.5:
                st.info("Similar Performance")
            elif diff > 0:
                st.success(f"{selected_names[0]} leads by {diff:.2f} points")
            else:
                st.warning(f"{selected_names[1]} leads by {abs(diff):.2f} points")
        
        with col3:
            st.markdown(f"#### 👤 {selected_names[1]}")
            st.write(f"**Email:** {candidate2_data['email']}")
            st.write(f"**Avg Rating:** {candidate2_data['avg_rating']:.2f}/10")
            st.write(f"**Sentiment:** {candidate2_data['sentiment']}")
            st.write(f"**Recommendation:** {candidate2_data['recommendation']}")
        
        st.markdown("---")
        
        # Side-by-side Radar Charts
        st.markdown("#### 🎯 Skills Comparison")
        col1, col2 = st.columns(2)
        
        categories = ['Technical', 'Communication', 'Coding', 'Problem Solving']
        
        with col1:
            values1 = [
                candidate1_data['technical_rating'],
                candidate1_data['communication_rating'],
                candidate1_data['coding_rating'],
                candidate1_data['problem_solving_rating']
            ]
            
            fig_radar1 = go.Figure()
            fig_radar1.add_trace(go.Scatterpolar(
                r=values1,
                theta=categories,
                fill='toself',
                name=selected_names[0],
                line=dict(color='#1f77b4', width=2),
                fillcolor='rgba(31, 119, 180, 0.3)'
            ))
            fig_radar1.update_layout(
                polar=dict(radialaxis=dict(visible=True, range=[0, 10])),
                showlegend=False,
                title=selected_names[0],
                height=400
            )
            st.plotly_chart(fig_radar1, use_container_width=True)
        
        with col2:
            values2 = [
                candidate2_data['technical_rating'],
                candidate2_data['communication_rating'],
                candidate2_data['coding_rating'],
                candidate2_data['problem_solving_rating']
            ]
            
            fig_radar2 = go.Figure()
            fig_radar2.add_trace(go.Scatterpolar(
                r=values2,
                theta=categories,
                fill='toself',
                name=selected_names[1],
                line=dict(color='#ff7f0e', width=2),
                fillcolor='rgba(255, 127, 14, 0.3)'
            ))
            fig_radar2.update_layout(
                polar=dict(radialaxis=dict(visible=True, range=[0, 10])),
                showlegend=False,
                title=selected_names[1],
                height=400
            )
            st.plotly_chart(fig_radar2, use_container_width=True)
        
        # Combined Comparison Chart
        st.markdown("#### 📊 Category-wise Comparison")
        comparison_data = pd.DataFrame({
            'Category': categories + categories,
            'Rating': values1 + values2,
            'Candidate': [selected_names[0]]*4 + [selected_names[1]]*4
        })
        
        fig_comparison = px.bar(
            comparison_data,
            x='Category',
            y='Rating',
            color='Candidate',
            barmode='group',
            text='Rating',
            color_discrete_map={selected_names[0]: '#1f77b4', selected_names[1]: '#ff7f0e'},
            range_y=[0, 10]
        )
        fig_comparison.update_traces(texttemplate='%{text:.1f}', textposition='outside')
        st.plotly_chart(fig_comparison, use_container_width=True)
        
        st.markdown("---")
    
    # Main Charts
    st.markdown("### 📊 Performance Analytics")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### ⭐ Average Ratings by Category")
        avg_ratings = {
            "Technical": filtered_df["technical_rating"].mean(),
            "Communication": filtered_df["communication_rating"].mean(),
            "Coding": filtered_df["coding_rating"].mean(),
            "Problem Solving": filtered_df["problem_solving_rating"].mean()
        }
        
        fig_pie = px.pie(
            values=list(avg_ratings.values()),
            names=list(avg_ratings.keys()),
            title="Distribution of Average Ratings",
            color_discrete_sequence=px.colors.sequential.RdBu,
            hole=0.4
        )
        fig_pie.update_traces(textposition='inside', textinfo='percent+label+value', 
                             texttemplate='%{label}<br>%{value:.1f}<br>(%{percent})')
        st.plotly_chart(fig_pie, use_container_width=True)
    
    with col2:
        st.markdown("#### 😊 Sentiment Count Distribution")
        sent_counts = filtered_df["sentiment"].value_counts()
        
        fig_sent_bar = px.bar(
            x=sent_counts.index,
            y=sent_counts.values,
            title="Sentiment Counts",
            labels={"x": "Sentiment", "y": "Count"},
            color=sent_counts.index,
            color_discrete_map={"Positive": "#28a745", "Neutral": "#ffc107", "Negative": "#dc3545"},
            text=sent_counts.values
        )
        fig_sent_bar.update_traces(texttemplate='%{text}', textposition='outside')
        fig_sent_bar.update_layout(showlegend=False)
        st.plotly_chart(fig_sent_bar, use_container_width=True)
    
    st.markdown("---")
    
    # Data Table
    st.markdown("### 📑 Filtered Interview Data")
    
    display_cols = ['name', 'email', 'technical_rating', 'communication_rating', 
                   'coding_rating', 'problem_solving_rating', 'avg_rating',
                   'sentiment', 'recommendation', 'timestamp']
    
    display_data = filtered_df[display_cols].copy()
    display_data['avg_rating'] = display_data['avg_rating'].round(2)
    display_data = display_data.sort_values('avg_rating', ascending=False)
    
    st.dataframe(display_data, use_container_width=True, hide_index=True)
    
    # Download Report Button
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col2:
        csv_data = filtered_df.to_csv(index=False)
        st.download_button(
            label="📥 Download Filtered Report (CSV)",
            data=csv_data,
            file_name=f"interview_analytics_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
            type="primary"
        )

# --------------------
# Main App with Tabs
# --------------------
def main():
    # Sidebar
    st.sidebar.title("🗂️ Navigation")
    
    st.sidebar.markdown("---")
    st.sidebar.header("Database Status")
    st.sidebar.metric("📊 Resumes in DB", collection.count())
    
    if os.path.exists("interview_ratings.csv"):
        ratings_df = pd.read_csv("interview_ratings.csv")
        ratings_df['recommendation'] = ratings_df['recommendation'].apply(clean_recommendation)
        st.sidebar.metric("⭐ Interviews Rated", len(ratings_df))
    
    st.sidebar.markdown("---")
    
    # Main Tabs
    tab1, tab2, tab3, tab4 = st.tabs([
        "📋 Job Matching",
        "📧 Email Scheduling",
        "⭐ Interview Ratings",
        "📊 Analytics Dashboard"
    ])
    
    # ==================== TAB 1: Job Matching ====================
    with tab1:
        st.title("🤖 AI-Powered Recruitment System")
        
        st.subheader("📍 Step 1 – Upload or Fetch Resumes")
        with st.expander("📂 Upload / Fetch Resumes", expanded=not st.session_state.step1_done):
            option = st.radio("Select Resume Input:", ["Upload Manually", "Fetch from Google Drive"])
            if option == "Upload Manually":
                files = st.file_uploader("Upload resumes (PDF)", type=["pdf"], accept_multiple_files=True)
                if files:
                    for f in files:
                        text = extract_text_from_pdf(f)
                        st.success(store_resume_text(text, f.name))
                    st.session_state.step1_done = True
            else:
                fid = st.text_input("Enter Google Drive Folder ID:")
                if st.button("📥 Fetch Resumes"):
                    with st.spinner("Fetching from Drive..."):
                        service = connect_drive()
                        results = fetch_resumes_from_drive(service, fid)
                        for r in results:
                            st.success(f"{r['filename']}: {r['store_msg']}")
                        st.session_state.step1_done = True
        
        if st.session_state.step1_done:
            st.markdown("---")
            st.subheader("🧾 Step 2 – Job Description Matching")
            jd = st.text_area("Paste Job Description:", height=200, value=st.session_state.jd_text)
            percentage = st.slider("Select Top Candidates by %:", 10, 100, 50, step=10)
            
            if st.button("🔍 Run Matching"):
                if jd.strip():
                    with st.spinner("Analyzing resumes..."):
                        explanation, results = find_best_match_with_llm(jd, percentage)
                    if results:
                        st.session_state.step2_done = True
                        st.session_state.jd_text = jd
                        st.success("✅ Matching Completed!")
                        st.markdown("### 🏆 Best Match Analysis")
                        st.markdown(explanation)
                else:
                    st.error("❗ Please enter a job description")
        
        if st.session_state.step2_done and os.path.exists("matched_candidates.csv"):
            st.markdown("---")
            st.subheader("📊 Step 3 – Matched Candidates")
            df = pd.read_csv("matched_candidates.csv")
            st.metric("📈 Candidates Selected", len(df))
            
            st.dataframe(df[["name", "email", "cosine_score", "llm_confidence"]], 
                        use_container_width=True)
            
            st.markdown("---")
            st.subheader("📝 Step 4 – Generate Interview Questions")
            st.markdown("Generate personalized interview questions for shortlisted candidates")
            
            selected_for_questions = st.selectbox(
                "Select Candidate for Interview Questions:",
                df["name"].tolist(),
                key="questions_candidate_select"
            )
            
            if st.button("🎯 Generate Interview Questions", type="primary"):
                candidate_row = df[df["name"] == selected_for_questions].iloc[0]
                
                with st.spinner(f"🤖 Generating personalized interview questions for {selected_for_questions}..."):
                    filename, filepath, questions_text = generate_interview_questions(
                        job_description=st.session_state.jd_text,
                        resume_text=candidate_row["resume_text"],
                        candidate_name=selected_for_questions
                    )
                    
                    if filename and filepath:
                        st.success(f"✅ Interview questions generated successfully!")
                        
                        df.loc[df["name"] == selected_for_questions, "interview_questions_doc"] = filename
                        df.to_csv("matched_candidates.csv", index=False)
                        
                        st.session_state.questions_generated[selected_for_questions] = {
                            "filename": filename,
                            "filepath": filepath
                        }
                        
                        with open(filepath, "rb") as file:
                            st.download_button(
                                label="📥 Download Interview Questions (Word Document)",
                                data=file,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_main_{selected_for_questions}"
                            )
                    else:
                        st.error(f"⚠️ Failed to generate questions: {questions_text}")
    
    # ==================== TAB 2: Email Scheduling ====================
    with tab2:
        st.title("📧 Interview Email Scheduling")
        
        if not st.session_state.step2_done or not os.path.exists("matched_candidates.csv"):
            st.warning("⚠️ Please complete Job Matching first (Tab 1)")
        else:
            df = pd.read_csv("matched_candidates.csv")
            
            st.subheader("📧 Step 5 – Send Interview Emails")
            
            selected_candidate = st.selectbox("Select Candidate:", df["name"].tolist(), key="email_candidate_select")
            row = df[df["name"] == selected_candidate].iloc[0]
            email = row["email"]
            
            st.write(f"📨 **Candidate Email:** {email}")
            st.write(f"📄 **Cosine Score:** {row['cosine_score']}")
            st.write(f"🤖 **LLM Confidence:** {row['llm_confidence']}")
            
            st.markdown("---")
            
            with st.expander("📋 View Candidate Disclaimer", expanded=False):
                st.markdown(row["disclaimer"])
            
            st.markdown("---")
        
            
            st.markdown("---")
            st.subheader("🚀 Automatic Scheduling")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("✉️ Send Shortlist Email", type="primary", key="send_shortlist"):
                    jd_text = st.session_state.get("jd_text", "")
                    if not jd_text:
                        st.error("❗ Please provide a Job Description in Tab 1 before sending.")
                    else:
                        res = send_shortlist_email(selected_candidate, email, jd_text)
                        st.success(res)
                        st.info("✅ Email sent! Candidate will reply with their availability.")
                        st.info("💡 Click the 'Check for Candidate Reply' button after candidate responds.")
            
            with col2:
                if st.button("🔍 Check for Candidate Reply", type="secondary", key="check_reply"):
                    with st.spinner(f"🔄 Scanning Gmail for reply from {email}..."):
                        interview_dt, msg = read_gmail_replies_and_extract_date(email)
                        
                        if interview_dt:
                            st.success(f"✅ Found reply! {msg}")
                            
                            send_result = send_interview_email(
                                name=selected_candidate,
                                email=email,
                                interview_datetime=interview_dt,
                                time_range_msg=msg
                            )
                            
                            st.success(send_result)
                            st.balloons()
                            
                            # Store in session state
                            st.session_state.scheduled_interviews[selected_candidate] = {
                                "datetime": interview_dt,
                                "email": email,
                                "msg": msg
                            }
                            
                            st.markdown("---")
                            st.markdown("### ✅ Interview Scheduled Successfully!")
                            st.write(f"**Candidate:** {selected_candidate}")
                            st.write(f"**Date & Time:** {interview_dt.strftime('%d %B %Y, %I:%M %p IST')}")
                            st.write(f"**Meeting Link:** https://meet.google.com/zag-iocp-kuz")
                            
                            calendar_link = create_google_calendar_link(selected_candidate, interview_dt)
                            st.markdown(f"[📅 Add to Your Calendar]({calendar_link})")
                            
                        else:
                            st.warning(f"⚠️ {msg}")
                            st.info("💡 Make sure the candidate has replied to your email with their availability.")
            
            # ✅ Manual Scheduling Section
            st.markdown("---")
            st.subheader("📅 Manual Interview Scheduling")
            st.markdown("If candidate provided availability outside email (phone/WhatsApp), schedule manually:")
            
            col_a, col_b = st.columns(2)
            
            with col_a:
                manual_date = st.date_input(
                    "Select Interview Date:", 
                    min_value=datetime.now().date(),
                    key="manual_date"
                )
            
            with col_b:
                manual_time = st.time_input(
                    "Select Interview Time:", 
                    value=datetime.strptime("10:00", "%H:%M").time(),
                    key="manual_time"
                )
            
            if st.button("📅 Schedule Interview Manually", type="secondary", key="manual_schedule"):
                manual_datetime = datetime.combine(manual_date, manual_time)
                
                send_result = send_interview_email(
                    name=selected_candidate,
                    email=email,
                    interview_datetime=manual_datetime,
                    time_range_msg="Manually scheduled by HR"
                )
                
                st.success(send_result)
                st.success(f"✅ Interview scheduled for {manual_datetime.strftime('%d %B %Y, %I:%M %p IST')}")
                
                # Store in session state
                st.session_state.scheduled_interviews[selected_candidate] = {
                    "datetime": manual_datetime,
                    "email": email,
                    "msg": "Manually scheduled"
                }
                
                st.balloons()
                
                calendar_link = create_google_calendar_link(selected_candidate, manual_datetime)
                st.markdown(f"[📅 Add to Your Calendar]({calendar_link})")
            
            # ✅ Display Scheduled Interviews
            if st.session_state.scheduled_interviews:
                st.markdown("---")
                st.subheader("📆 Scheduled Interviews")
                
                scheduled_data = []
                for name, info in st.session_state.scheduled_interviews.items():
                    scheduled_data.append({
                        "Candidate": name,
                        "Email": info["email"],
                        "Date & Time": info["datetime"].strftime('%d %B %Y, %I:%M %p IST'),
                        "Status": "✅ Confirmed"
                    })
                
                if scheduled_data:
                    scheduled_df = pd.DataFrame(scheduled_data)
                    st.dataframe(scheduled_df, use_container_width=True, hide_index=True)
    
    # ==================== TAB 3: Interview Ratings ====================
    with tab3:
        st.title("⭐ Interview Rating & Feedback")
        
        st.markdown("### 📝 Rate Interview Performance")
        st.markdown("Rate the candidate on the following parameters (1-10):")
        
        with st.form("rating_form", clear_on_submit=False):
            col1, col2 = st.columns(2)
            
            with col1:
                candidate_name_rating = st.text_input("Candidate Name:", key="rating_name")
                candidate_email_rating = st.text_input("Candidate Email:", key="rating_email")
                technical_rating = st.slider("🔧 Technical Skills", 1, 10, 5, key="tech")
                communication_rating = st.slider("💬 Communication", 1, 10, 5, key="comm")
            
            with col2:
                st.write("")
                st.write("")
                coding_rating = st.slider("💻 Coding Ability", 1, 10, 5, key="code")
                problem_solving_rating = st.slider("🧩 Problem Solving", 1, 10, 5, key="prob")
            
            interviewer_notes = st.text_area(
                "📋 Interviewer Notes & Comments:", 
                height=150,
                placeholder="Add detailed observations about the candidate's performance, specific examples from the interview, strengths demonstrated, areas of concern, cultural fit indicators, and any other relevant insights that informed your ratings...",
                key="notes"
            )
            
            submit_button = st.form_submit_button("🎯 Submit Rating & Generate AI Analysis", type="primary")
        
        if submit_button:
            if candidate_name_rating and candidate_email_rating:
                with st.spinner("🤖 Analyzing feedback with advanced AI evaluation system..."):
                    llm_feedback = analyze_interview_feedback(
                        technical_rating, 
                        communication_rating, 
                        coding_rating,
                        problem_solving_rating,
                        interviewer_notes
                    )
                    
                    rating_data = {
                        "name": candidate_name_rating,
                        "email": candidate_email_rating,
                        "technical_rating": technical_rating,
                        "communication_rating": communication_rating,
                        "coding_rating": coding_rating,
                        "problem_solving_rating": problem_solving_rating,
                        "interviewer_notes": interviewer_notes,
                        "sentiment": llm_feedback["sentiment"],
                        "recommendation": llm_feedback["recommendation"],
                        "llm_analysis": llm_feedback["analysis"],
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    
                    store_interview_rating(rating_data)
                    
                    st.session_state.last_rating_result = {
                        "rating_data": rating_data,
                        "llm_feedback": llm_feedback,
                        "ratings": {
                            "technical": technical_rating,
                            "communication": communication_rating,
                            "coding": coding_rating,
                            "problem_solving": problem_solving_rating
                        }
                    }
                    
                    st.success("✅ Interview rating saved successfully to database!")
                    
            else:
                st.error("❗ Please enter both candidate name and email address")
        
        if st.session_state.last_rating_result:
            result = st.session_state.last_rating_result
            llm_feedback = result["llm_feedback"]
            rating_data = result["rating_data"]
            
            st.markdown("---")
            st.markdown("### 🤖 AI-Powered Professional Analysis")
            st.markdown(f"**Candidate:** {rating_data['name']} | **Email:** {rating_data['email']}")
            
            col_a, col_b, col_c = st.columns(3)
            
            with col_a:
                sentiment_emoji = {
                    "Positive": "🟢",
                    "Neutral": "🟡",
                    "Negative": "🔴"
                }
                st.metric(
                    "📊 Sentiment", 
                    f"{sentiment_emoji.get(llm_feedback['sentiment'], '⚪')} {llm_feedback['sentiment']}"
                )
            
            with col_b:
                recommendation_emoji = {
                    "Hire": "✅",
                    "Review": "⚠️",
                    "Reject": "❌"
                }
                st.metric(
                    "🎯 Recommendation", 
                    f"{recommendation_emoji.get(llm_feedback['recommendation'], '⚪')} {llm_feedback['recommendation']}"
                )
            
            with col_c:
                avg_score = llm_feedback["avg_score"]
                score_color = "🟢" if avg_score >= 7 else "🟡" if avg_score >= 5 else "🔴"
                st.metric("⭐ Average Rating", f"{score_color} {avg_score:.1f} / 10")
            
            st.markdown("---")
            st.markdown("#### 📋 Comprehensive Professional Analysis")
            st.info(llm_feedback["analysis"])
        
        # View All Ratings
        st.markdown("---")
        st.subheader("📈 All Interview Ratings")
        
        if os.path.exists("interview_ratings.csv"):
            ratings_df = pd.read_csv("interview_ratings.csv")
            ratings_df['recommendation'] = ratings_df['recommendation'].apply(clean_recommendation)
            ratings_df['sentiment'] = ratings_df['sentiment'].apply(clean_sentiment)
            
            if 'timestamp' not in ratings_df.columns:
                ratings_df['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            if len(ratings_df) > 0:
                st.markdown(f"**Total Interviews Rated:** {len(ratings_df)}")
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    hire_count = len(ratings_df[ratings_df["recommendation"] == "Hire"])
                    st.metric("✅ Hire", hire_count)
                with col2:
                    review_count = len(ratings_df[ratings_df["recommendation"] == "Review"])
                    st.metric("⚠️ Review", review_count)
                with col3:
                    reject_count = len(ratings_df[ratings_df["recommendation"] == "Reject"])
                    st.metric("❌ Reject", reject_count)
                with col4:
                    avg_technical = ratings_df["technical_rating"].mean()
                    st.metric("Avg Technical", f"{avg_technical:.1f}/10")
                
                display_df = ratings_df[[
                    "name", "email", "technical_rating", "communication_rating", 
                    "coding_rating", "problem_solving_rating",
                    "sentiment", "recommendation", "timestamp"
                ]].copy()
                
                display_df["avg_rating"] = (
                    display_df["technical_rating"] + 
                    display_df["communication_rating"] + 
                    display_df["coding_rating"] + 
                    display_df["problem_solving_rating"]
                ) / 4
                
                st.dataframe(display_df, use_container_width=True)
                
                with st.expander("🔍 View Detailed Feedback for Specific Candidate"):
                    selected_rating = st.selectbox(
                        "Select candidate to view detailed analysis:",
                        ratings_df["name"].tolist()
                    )
                    
                    rating_row = ratings_df[ratings_df["name"] == selected_rating].iloc[0]
                    
                    st.markdown(f"### 👤 Candidate: {rating_row['name']}")
                    st.markdown(f"**📧 Email:** {rating_row['email']}")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("🔧 Technical Skills", f"{rating_row['technical_rating']}/10")
                        st.metric("💬 Communication", f"{rating_row['communication_rating']}/10")
                    with col2:
                        st.metric("💻 Coding Ability", f"{rating_row['coding_rating']}/10")
                        st.metric("🧩 Problem Solving", f"{rating_row['problem_solving_rating']}/10")
                    
                    st.markdown("---")
                    st.markdown("#### 📝 Interviewer's Original Notes:")
                    st.write(rating_row['interviewer_notes'])
                    
            else:
                st.info("📭 No interview ratings yet. Complete some interviews to see data here.")
        else:
            st.info("📭 No interview ratings database found. Rate your first interview to get started!")
    
    # ==================== TAB 4: Analytics Dashboard ====================
    with tab4:
        create_analytics_dashboard()

if __name__ == "__main__":
    main()